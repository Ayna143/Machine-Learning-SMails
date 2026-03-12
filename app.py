"""
Flask web application for Email Spam Detection.
Features: single check, batch import, image OCR, model comparison dashboard.
"""

import os
import io
import csv
import json
import traceback
from flask import Flask, render_template, request, jsonify
import pandas as pd

from src.predict import SpamDetector

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

detector = None
MODELS_DIR = 'models'


def get_detector():
    global detector
    if detector is None:
        detector = SpamDetector(model_dir=MODELS_DIR)
    return detector


# ── Pages ────────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    models = get_detector().get_available_models()
    return render_template('index.html', models=models)


# ── Single Email Prediction ─────────────────────────────────────────────────

@app.route('/predict', methods=['POST'])
def predict():
    try:
        data = request.get_json()
        email_text = data.get('email_text', '').strip()
        sender = data.get('sender', '').strip()
        device = data.get('device', '').strip()
        model_name = data.get('model_name', None)

        if not email_text:
            return jsonify({'error': 'Please enter email text to analyze.'}), 400

        result = get_detector().predict(email_text, sender, device, model_name)
        return jsonify(result)

    except FileNotFoundError as e:
        return jsonify({'error': str(e)}), 500
    except Exception:
        traceback.print_exc()
        return jsonify({'error': 'An unexpected error occurred.'}), 500


# ── Batch / CSV Import ──────────────────────────────────────────────────────

@app.route('/predict_batch', methods=['POST'])
def predict_batch():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded.'}), 400

        file = request.files['file']
        model_name = request.form.get('model_name', None)

        if file.filename == '':
            return jsonify({'error': 'Empty filename.'}), 400
        if not file.filename.lower().endswith(('.csv', '.txt', '.tsv')):
            return jsonify({'error': 'Please upload a CSV or TXT file.'}), 400

        stream = io.StringIO(file.stream.read().decode('utf-8', errors='replace'))
        sample = stream.read(4096)
        stream.seek(0)
        dialect = csv.Sniffer().sniff(sample, delimiters=',\t|;')
        df = pd.read_csv(stream, sep=dialect.delimiter)

        text_col = None
        for c in ['text', 'message', 'email', 'content', 'Message', 'Text', 'v2']:
            if c in df.columns:
                text_col = c
                break
        if text_col is None:
            text_col = df.columns[-1]

        label_col = None
        for c in ['label', 'Label', 'spam', 'Spam', 'class', 'v1']:
            if c in df.columns:
                label_col = c
                break

        texts = df[text_col].fillna('').astype(str).tolist()
        results = get_detector().predict_batch(texts, model_name=model_name)

        spam_count = sum(1 for r in results if r['is_spam'])
        ham_count = len(results) - spam_count

        accuracy = None
        if label_col is not None:
            true_labels = df[label_col].tolist()
            mapped = []
            for lbl in true_labels:
                lbl_str = str(lbl).strip().lower()
                mapped.append(1 if lbl_str in ('1', 'spam') else 0)
            predicted = [1 if r['is_spam'] else 0 for r in results]
            correct = sum(1 for t, p in zip(mapped, predicted) if t == p)
            accuracy = round(correct / max(len(mapped), 1) * 100, 2)

        preview = []
        for i, r in enumerate(results[:100]):
            preview.append({
                'index': i + 1,
                'text': texts[i][:120] + ('...' if len(texts[i]) > 120 else ''),
                'prediction': r['prediction'],
                'confidence': r['confidence'],
            })

        return jsonify({
            'total': len(results),
            'spam_count': spam_count,
            'ham_count': ham_count,
            'accuracy': accuracy,
            'preview': preview,
        })

    except Exception:
        traceback.print_exc()
        return jsonify({'error': 'Failed to process the uploaded file.'}), 500


# ── Document (PDF / Word) Prediction ────────────────────────────────────────

def _extract_text_from_document(file_stream, filename):
    """Extract text from PDF or Word (.docx). Returns (text, error_message)."""
    ext = (filename or '').lower().split('.')[-1]
    if ext == 'pdf':
        try:
            import fitz  # pymupdf
        except ImportError:
            return ('', 'PDF support requires: pip install pymupdf')
        try:
            doc = fitz.open(stream=file_stream.read(), filetype='pdf')
            parts = []
            for page in doc:
                parts.append(page.get_text())
            doc.close()
            return (' '.join(parts).strip(), None)
        except Exception as e:
            return ('', f'PDF extraction failed: {e}')
    if ext in ('docx', 'doc'):
        try:
            from docx import Document
        except ImportError:
            return ('', 'Word support requires: pip install python-docx')
        try:
            doc = Document(io.BytesIO(file_stream.read()))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return (' '.join(parts).strip(), None)
        except Exception as e:
            return ('', f'Word extraction failed: {e}')
    return ('', 'Unsupported format. Use PDF or Word (.docx).')


@app.route('/predict_document', methods=['POST'])
def predict_document():
    try:
        if 'document' not in request.files:
            return jsonify({'error': 'No file uploaded.'}), 400

        file = request.files['document']
        model_name = request.form.get('model_name', None)

        if file.filename == '':
            return jsonify({'error': 'Empty filename.'}), 400

        allowed_ext = ('.pdf', '.docx', '.doc')
        if not file.filename.lower().endswith(allowed_ext):
            return jsonify({'error': 'Please upload a PDF or Word (.docx) file.'}), 400

        file.stream.seek(0)
        extracted_text, err = _extract_text_from_document(file.stream, file.filename)
        if err:
            return jsonify({'error': err}), 500

        if not extracted_text:
            return jsonify({
                'extracted_text': '',
                'prediction': 'not spam',
                'is_spam': False,
                'confidence': None,
                'reasons': ['No text could be extracted from the document.'],
                'features': {},
                'model_used': model_name or 'best',
            })

        result = get_detector().predict(extracted_text, model_name=model_name)
        result['extracted_text'] = extracted_text
        return jsonify(result)

    except Exception:
        traceback.print_exc()
        return jsonify({'error': 'Failed to process the document.'}), 500


# ── Image OCR Prediction ────────────────────────────────────────────────────

@app.route('/predict_image', methods=['POST'])
def predict_image():
    try:
        if 'image' not in request.files:
            return jsonify({'error': 'No image uploaded.'}), 400

        file = request.files['image']
        model_name = request.form.get('model_name', None)

        if file.filename == '':
            return jsonify({'error': 'Empty filename.'}), 400

        allowed_ext = ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp')
        if not file.filename.lower().endswith(allowed_ext):
            return jsonify({'error': 'Please upload an image file (PNG, JPG, etc).'}), 400

        try:
            from PIL import Image
            import numpy as np
        except ImportError:
            return jsonify({'error': 'Pillow or numpy is not installed. Run: pip install Pillow numpy'}), 500

        # Read file into memory so stream is not exhausted
        try:
            raw = file.read()
            if not raw:
                return jsonify({'error': 'Image file is empty.'}), 400
            image = Image.open(io.BytesIO(raw)).convert('RGB')
        except Exception:
            traceback.print_exc()
            return jsonify({'error': 'Invalid or unsupported image. Try a different file (PNG/JPG).'}), 400

        img_array = np.array(image)

        extracted_text = ''
        try:
            import pytesseract
            extracted_text = pytesseract.image_to_string(image).strip()
        except Exception:
            pass

        if not extracted_text or not extracted_text.strip():
            try:
                import easyocr
                if not hasattr(predict_image, '_easyocr_reader'):
                    predict_image._easyocr_reader = easyocr.Reader(['en'], gpu=False, verbose=False)
                reader = predict_image._easyocr_reader
                results = reader.readtext(img_array)
                extracted_text = ' '.join([r[1] for r in results]).strip()
            except ImportError:
                return jsonify({
                    'error': 'OCR not available. Install Tesseract (add to PATH) or run: pip install easyocr'
                }), 500
            except Exception:
                traceback.print_exc()
                return jsonify({
                    'error': 'OCR failed. Try a clearer image with readable text, or install Tesseract and add it to PATH.'
                }), 500

        if not extracted_text.strip():
            return jsonify({
                'extracted_text': '',
                'prediction': 'not spam',
                'is_spam': False,
                'confidence': None,
                'reasons': ['No text could be extracted from the image.'],
                'features': {},
                'model_used': model_name or 'best',
            })

        try:
            result = get_detector().predict(extracted_text, model_name=model_name)
        except Exception:
            traceback.print_exc()
            return jsonify({'error': 'Spam check failed. Make sure models are trained (run train_and_evaluate.py).'}), 500

        result['extracted_text'] = extracted_text.strip()
        return jsonify(result)

    except Exception:
        traceback.print_exc()
        return jsonify({'error': 'Failed to process the image. Check the terminal/console where the server is running for the error details.'}), 500


# ── Model Comparison Data ───────────────────────────────────────────────────

@app.route('/comparison')
def comparison():
    metrics_path = os.path.join(MODELS_DIR, 'metrics.json')
    if not os.path.exists(metrics_path):
        return jsonify({'error': 'No metrics found. Run training first.'}), 404

    with open(metrics_path) as f:
        metrics = json.load(f)
    return jsonify(metrics)


@app.route('/models_list')
def models_list():
    return jsonify({'models': get_detector().get_available_models()})


# ── Run ──────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print("\n  Starting Email Spam Detection Web App...")
    print("  Open http://127.0.0.1:5000 in your browser.\n")
    app.run(debug=True, host='127.0.0.1', port=5000, use_reloader=False)
