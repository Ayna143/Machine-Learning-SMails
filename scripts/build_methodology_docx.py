from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def _font_run(run, *, bold: bool = False):
    run.font.name = "Poppins"
    run.font.size = Pt(12)
    run.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Poppins")
    rFonts.set(qn("w:hAnsi"), "Poppins")

def add_para(doc: Document, text: str, *, bold: bool = False, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    _font_run(run, bold=bold)
    return p

def main():
    root = Path(__file__).resolve().parents[1]
    out_dir = root / "docuuu"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Data_Preprocessing_and_Feature_Engineering.docx"

    doc = Document()

    add_para(
        doc,
        "3.5 Data Preprocessing",
        bold=True,
    )
    add_para(
        doc,
        "Raw email text was cleaned and prepared before model training to ensure consistency and "
        "improve data quality. Because the production pipeline combines traditional bag-of-words "
        "features with dense semantic embeddings, preprocessing was organized into paths that serve "
        "each representation appropriately.",
    )

    subs_35 = [
        (
            "Text Normalization for TF-IDF.",
            "For the lexical vectorization branch, all text was converted to lowercase to standardize "
            "word representation. Non-alphabetic characters such as punctuation and symbols were removed "
            "using the regular expression [^a-zA-Z\\s]. This ensured that superficial variants such as "
            "“FREE” and “free” were aligned before tokenization.",
        ),
        (
            "Tokenization and Stop Word Removal.",
            "For the TF-IDF branch only, text was split into individual words using the NLTK library. "
            "Common English stop words (e.g., “the,” “is,” “for”) and tokens shorter than two characters "
            "were removed, since they carry little discriminative value for spam identification. The "
            "retained tokens were rejoined into a single cleaned string for vectorization.",
        ),
        (
            "Embedding-Oriented Text Preparation.",
            "For the semantic embedding branch, the system did not strip punctuation or apply stop-word "
            "removal. Instead, raw message text underwent light normalization—whitespace was collapsed "
            "and an upper length bound was applied to limit memory use—so that wording, capitalization, "
            "and symbols that may carry phishing or promotional cues remained available to the encoder. "
            "This complements the aggressively cleaned TF-IDF input by preserving contextual meaning.",
        ),
        (
            "Handling Missing Values.",
            "Rows with missing text or label values were dropped. Optional fields (sender, device) were "
            "filled with empty strings. Non-string values were converted to strings to prevent errors "
            "during feature extraction and vectorization.",
        ),
    ]
    for title, body in subs_35:
        p = doc.add_paragraph()
        r1 = p.add_run(title + " ")
        _font_run(r1, bold=True)
        r2 = p.add_run(body)
        _font_run(r2)

    add_para(doc, "TF-IDF Vectorization.", bold=True)
    add_para(
        doc,
        "The cleaned token strings produced for TF-IDF were converted into numerical vectors using "
        "Scikit-learn’s TfidfVectorizer with the following parameters:",
    )
    for bullet in (
        "max_features = 3,000",
        "ngram_range = (1, 2) — unigrams and bigrams",
        "sublinear_tf = True — log normalization applied to term frequencies",
    ):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bullet)
        _font_run(run)

    add_para(
        doc,
        "This configuration captured single-word and two-word patterns (e.g., “free,” “click here”) "
        "while capping vocabulary size. The output was a sparse matrix of TF-IDF weights per email.",
    )

    add_para(doc, "Dense Semantic Embedding.", bold=True)
    p = doc.add_paragraph()
    t = (
        "In parallel, minimally prepared text was encoded with a pretrained sentence-transformer model "
        "(default: sentence-transformers/all-MiniLM-L6-v2), producing one dense vector per email. "
        "Embedding dimension depends on the chosen model (384 dimensions for this default). For "
        "integration with multinomial-style components in the hybrid stack, semantic dimensions were "
        "clipped at zero so all combined dense blocks were non-negative where required."
    )
    r = p.add_run(t)
    _font_run(r)

    add_para(
        doc,
        "Train-Test Split.",
        bold=True,
    )
    p = doc.add_paragraph()
    r = p.add_run(
        "The dataset was split 80% for training and 20% for testing with stratified sampling. "
        "Three-fold cross-validation was also applied under the training configuration."
    )
    _font_run(r)

    doc.add_paragraph()

    add_para(doc, "3.6 Feature Engineering", bold=True)
    add_para(
        doc,
        "In addition to TF-IDF and semantic embedding features, a total of sixteen (16) engineered "
        "features were extracted to capture content, structure, domain, sender, and device-related "
        "attributes, following feature engineering approaches common in spam detection research.",
    )

    rows = [
        ("Feature", "Category", "Description"),
        (
            "suspicious_keyword_count",
            "content",
            "Weighted score from known spam trigger phrases (derived from raw hit counts, scaled and "
            "capped so classification relies jointly on lexical, semantic, and structural signals)",
        ),
        ("url_count", "content", "Number of URLs detected using regex (http/https/www patterns)"),
        (
            "special_char_count",
            "content",
            "Count of characters such as !, $, %, &, *, #, @, ^, ~",
        ),
        (
            "uppercase_ratio",
            "content",
            "Proportion of uppercase characters to total characters",
        ),
        ("email_length", "Structural", "Total word count of the email"),
        (
            "has_html",
            "Structural",
            "Binary flag: 1 if HTML tags are detected in the text",
        ),
        (
            "has_suspicious_domain",
            "Domain",
            "Binary flag: 1 if an @domain segment uses a suspicious TLD, matches a subdomain "
            "heuristic such as digits in the hostname segment, or triggers other extractor rules",
        ),
        ("exclamation_count", "content", "Total number of exclamation marks"),
        ("dollar_sign_count", "content", "Total number of dollar sign characters"),
        (
            "digit_ratio",
            "content",
            "Proportion of digit characters to total characters",
        ),
        (
            "sender_domain_suspicious",
            "Sender",
            "1 if the sender domain uses a suspicious TLD or matches lookalike brand patterns",
        ),
        (
            "sender_has_numbers",
            "Sender",
            "1 if the sender’s local-part contains numeric characters",
        ),
        (
            "sender_domain_length",
            "Sender",
            "Length of the sender’s domain string (capped at 50)",
        ),
        (
            "device_is_unknown",
            "Device",
            "1 if the sending device/client is unknown or absent",
        ),
        (
            "device_is_mobile",
            "Device",
            "1 if the sending device matches known mobile identifiers",
        ),
        (
            "device_is_automated",
            "Device",
            "1 if the sending device matches bulk mail or automated-system keywords",
        ),
    ]

    table = doc.add_table(rows=len(rows), cols=3)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        cells = table.rows[i].cells
        for j, cell_text in enumerate(row):
            cells[j].text = ""
            para = cells[j].paragraphs[0]
            run = para.add_run(cell_text)
            _font_run(run, bold=(i == 0))

    add_para(doc, "Table 5: Feature Engineering", bold=True)

    add_para(
        doc,
        "Suspicious TLDs flagged by the system included .xyz, .top, .click, .info, .bid, .tk, .ml, "
        ".ga, and .cf. Lookalike domain substring checks included patterns such as paypa1, amaz0n, "
        "g00gle, app1e, and micros0ft within the sender domain.",
    )
    add_para(
        doc,
        "The engineered features were assembled as a dense NumPy array and concatenated horizontally "
        "with the sparse TF-IDF block and the sparse (clipped, non-negative) embedding block using "
        "scipy.sparse.hstack prior to model training. Under the default hybrid configuration, the "
        "combined dimensionality equals 3,000 TF-IDF features plus the embedding width d plus 16 "
        "engineered features (e.g., 3,000 + 384 + 16 = 3,400 for the default MiniLM encoder). No "
        "dimensionality reduction such as PCA was applied, as the resulting space remained "
        "manageable for the selected models.",
    )

    doc.save(out_path)
    print(f"Wrote {out_path}")

if __name__ == "__main__":
    main()
